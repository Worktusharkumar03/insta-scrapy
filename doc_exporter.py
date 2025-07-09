# Document Exporter Module
# Saves results into DOCX (and optionally PDF)

from docx import Document

def save_to_docx(metadata, original_text, rewritten_text, output_path):
    doc = Document()
    doc.add_heading('Instagram Video Transcript', 0)
    doc.add_heading('Metadata', level=1)
    for k, v in metadata.items():
        doc.add_paragraph(f'{k}: {v}')
    doc.add_heading('Original Transcript', level=1)
    doc.add_paragraph(original_text)
    doc.add_heading('Rewritten Transcript', level=1)
    doc.add_paragraph(rewritten_text)
    doc.save(output_path)

# Example usage:
# save_to_docx({'url': '...'}, 'original', 'rewritten', 'output.docx')
